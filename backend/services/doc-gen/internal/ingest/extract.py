#!/usr/bin/env python3
"""doc-gen 文本提取子进程（ADR-0012）。

被 Go 侧 ingest 通过 exec.Command 调用：
    python3 extract.py <文件路径> [格式]

格式为扩展名（不含点），缺省时从路径推断；Go 侧已规范化双扩展名
（如 .pdf.pdf -> pdf）。输出 stdout 一行 JSON：

    {"text": "...", "tables": ["| .. |"], "warnings": [".."], "needs_ocr": false}

按格式分流，依赖库缺失时自动降级，不整体失败。
"""
import json
import os
import re
import sys
import traceback


def rows_to_md(rows):
    """把二维单元格列表转为 Markdown 表格。"""
    if not rows:
        return ""
    width = max(len(r) for r in rows)
    rows = [list(r) + [""] * (width - len(r)) for r in rows]
    header = "| " + " | ".join(rows[0]) + " |"
    sep = "| " + " | ".join(["---"] * width) + " |"
    body = "\n".join("| " + " | ".join(r) + " |" for r in rows[1:])
    return "\n".join([header, sep, body])


def extract_docx(path):
    try:
        import docx
    except ImportError:
        return fallback_docx(path)
    parts = []
    tables = []
    doc = docx.Document(path)
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    for tbl in doc.tables:
        rows = [[c.text.strip().replace("\n", " ") for c in row.cells] for row in tbl.rows]
        md = rows_to_md(rows)
        if md:
            tables.append(md)
            parts.append(md)
    return {"text": "\n\n".join(parts), "tables": tables, "warnings": []}


def fallback_docx(path):
    import zipfile
    try:
        z = zipfile.ZipFile(path)
        data = z.read("word/document.xml").decode("utf-8", "ignore")
        text = re.sub(r"</w:p>", "\n", data)
        text = re.sub(r"<[^>]+>", "", text)
        return {"text": text, "tables": [], "warnings": ["docx: python-docx 缺失，降级 zip+XML"]}
    except Exception as e:
        return {"text": "", "tables": [], "warnings": ["docx 降级失败: " + str(e)]}


def extract_pdf(path):
    # 文本优先 pypdfium2（逐页 close，内存友好，适合大 PDF）
    parts = []
    try:
        import pypdfium2 as pdfium
        doc = pdfium.PdfDocument(path)
        for i in range(len(doc)):
            page = doc[i]
            tp = page.get_textpage()
            t = tp.get_text_range() or ""
            tp.close()
            page.close()
            if t.strip():
                parts.append(t)
        doc.close()
    except ImportError:
        return _extract_pdf_pdfplumber(path)
    except Exception as e:
        return {"text": "", "tables": [], "warnings": ["pdf 提取异常: " + str(e)], "needs_ocr": True}
    text = "\n\n".join(parts)
    needs_ocr = len(text.strip()) < 50
    # 表格仅对小文件用 pdfplumber（大文件避免 OOM）
    tables = []
    if not needs_ocr:
        try:
            size = os.path.getsize(path)
        except OSError:
            size = 0
        if size < 10 * 1024 * 1024:
            tables = _extract_pdf_tables(path)
    return {"text": text, "tables": tables, "warnings": [], "needs_ocr": needs_ocr}


def _extract_pdf_pdfplumber(path):
    try:
        import pdfplumber
    except ImportError:
        return {"text": "", "tables": [], "warnings": ["pdf: pdfplumber/pypdfium2 缺失"], "needs_ocr": True}
    parts = []
    tables = []
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text() or ""
                if t:
                    parts.append(t)
                for tbl in (page.extract_tables() or []):
                    rows = [[(c or "").replace("\n", " ") for c in r] for r in tbl]
                    md = rows_to_md(rows)
                    if md:
                        tables.append(md)
                        parts.append(md)
    except Exception as e:
        return {"text": "", "tables": [], "warnings": ["pdf 提取异常: " + str(e)], "needs_ocr": True}
    text = "\n\n".join(parts)
    return {"text": text, "tables": tables, "warnings": [], "needs_ocr": len(text.strip()) < 50}


def _extract_pdf_tables(path):
    tables = []
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                for tbl in (page.extract_tables() or []):
                    rows = [[(c or "").replace("\n", " ") for c in r] for r in tbl]
                    md = rows_to_md(rows)
                    if md:
                        tables.append(md)
    except Exception:
        pass
    return tables


def extract_xlsx(path):
    try:
        import openpyxl
    except ImportError:
        return fallback_xlsx(path)
    parts = []
    wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            cells = [("" if c is None else str(c)).strip() for c in row]
            if any(cells):
                parts.append(" | ".join(cells))
    return {"text": "\n".join(parts), "tables": [], "warnings": []}


def fallback_xlsx(path):
    import zipfile
    try:
        z = zipfile.ZipFile(path)
        data = z.read("xl/sharedStrings.xml").decode("utf-8", "ignore")
        text = re.sub(r"</si>", "\n", data)
        text = re.sub(r"<[^>]+>", "", text)
        return {"text": text, "tables": [], "warnings": ["xlsx: openpyxl 缺失，降级 sharedStrings"]}
    except Exception as e:
        return {"text": "", "tables": [], "warnings": ["xlsx 降级失败: " + str(e)]}


def extract_xls(path):
    return {"text": "", "tables": [], "warnings": ["xls: 需 libreoffice 或 xlrd<2 转换，已跳过"], "needs_ocr": True}


def extract_text_file(path):
    for enc in ("utf-8", "gbk", "gb2312", "latin-1"):
        try:
            with open(path, "r", encoding=enc) as f:
                return {"text": f.read(), "tables": [], "warnings": []}
        except (UnicodeDecodeError, LookupError):
            continue
    return {"text": "", "tables": [], "warnings": ["text: 未知编码"]}


EXTRACTORS = {
    "docx": extract_docx,
    "pdf": extract_pdf,
    "xlsx": extract_xlsx,
    "xls": extract_xls,
}

TEXT_EXTS = {"", ".txt", ".md", ".csv", ".json", ".yaml", ".yml", ".log", ".xml", ".html", ".htm"}


def main():
    if len(sys.argv) < 2:
        print(json.dumps({"text": "", "tables": [], "warnings": ["缺少路径参数"]}))
        return
    path = sys.argv[1]
    fmt = sys.argv[2].lower() if len(sys.argv) > 2 and sys.argv[2] else ""
    if not fmt:
        fmt = os.path.splitext(path)[1].lower().lstrip(".")
        if not fmt and path.lower().endswith(".pdf.pdf"):
            fmt = "pdf"

    fn = EXTRACTORS.get(fmt)
    if fn is None:
        res = extract_text_file(path)
    else:
        try:
            res = fn(path)
        except Exception as e:
            res = {"text": "", "tables": [], "warnings": [fmt + " 提取失败: " + str(e) + "\n" + traceback.format_exc()]}

    res.setdefault("text", "")
    res.setdefault("tables", [])
    res.setdefault("warnings", [])
    res.setdefault("needs_ocr", False)
    sys.stdout.write(json.dumps(res, ensure_ascii=False))
    sys.stdout.write("\n")


if __name__ == "__main__":
    main()
